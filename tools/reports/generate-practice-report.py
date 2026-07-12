from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "reports" / "practice-report-source.md"
OUTPUT = ROOT / "docs" / "reports" / "Отчёт_по_учебной_практике_ThermoSelect.docx"
EMBLEM = ROOT / "docs" / "reports" / "assets" / "institute-emblem.png"

BODY_FONT = "Times New Roman"
CODE_FONT = "Courier New"
BODY_SIZE = Pt(14)
USABLE_WIDTH_DXA = 9690


def set_run_font(run, name=BODY_FONT, size=BODY_SIZE, bold=None, italic=None):
    run.font.name = name
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, rows):
    col_count = len(rows[0])
    weights = []
    for col in range(col_count):
        longest = max(len(str(row[col])) for row in rows)
        weights.append(max(12, min(45, longest)))
    total = sum(weights)
    widths = [round(USABLE_WIDTH_DXA * weight / total) for weight in weights]
    widths[-1] += USABLE_WIDTH_DXA - sum(widths)

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill="E7E7E7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))
    set_run_font(run, size=Pt(12))


def add_toc(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = r' TOC \o "1-3" \h \z \u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Содержание будет обновлено при открытии документа в Microsoft Word."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, placeholder, end))
    set_run_font(run)


def set_update_fields(document):
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_inline(paragraph, text, size=BODY_SIZE):
    token_pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    position = 0
    for match in token_pattern.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position:match.start()]), size=size)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size=size, bold=True)
        else:
            set_run_font(paragraph.add_run(token[1:-1]), name=CODE_FONT, size=Pt(max(9, size.pt - 2)))
        position = match.end()
    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]), size=size)


def add_cover(document):
    section = document.sections[0]
    section.different_first_page_header_footer = True
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    emblem = paragraph.add_run().add_picture(str(EMBLEM), width=Inches(0.68))
    emblem._inline.docPr.set("descr", "Эмблема СПбГТИ(ТУ)")

    for text, bold, size in (
        ("МИНОБРНАУКИ РОССИИ", False, 14),
        ("федеральное государственное бюджетное образовательное учреждение высшего образования", False, 12),
        ("«Санкт-Петербургский государственный технологический институт (технический университет)»", False, 12),
        ("СПбГТИ(ТУ)", True, 12),
    ):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(text), size=Pt(size), bold=bold)

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    set_run_font(p.add_run("ОТЧЁТ"), size=Pt(16), bold=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    set_run_font(p.add_run("ПО УЧЕБНОЙ ПРАКТИКЕ"), size=Pt(16), bold=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    set_run_font(p.add_run("«Разработка информационно-поисковой системы для выбора теплообменных аппаратов»"), bold=True)

    document.add_paragraph()
    metadata = [
        ("Направление подготовки", "09.03.01 — Информатика и вычислительная техника"),
        ("Направленность", "Системы автоматизированного проектирования"),
        ("Факультет", "Информационных технологий и управления"),
        ("Кафедра", "Систем автоматизированного проектирования и управления"),
        ("Студент", "____________________________________________"),
        ("Группа", "____________"),
        ("Руководитель практики", "____________________________________________"),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in metadata:
        cells = table.add_row().cells
        for cell, text in zip(cells, (label, value)):
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.first_line_indent = Cm(0)
            add_inline(para, text, Pt(12))
    set_table_geometry(table, metadata)

    for _ in range(2):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    set_run_font(p.add_run("Санкт-Петербург"), size=Pt(12))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    set_run_font(p.add_run("2026"), size=Pt(12))
    document.add_page_break()


def add_assignment(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    set_run_font(p.add_run("ЗАДАНИЕ НА ТЕХНОЛОГИЧЕСКУЮ"), bold=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    set_run_font(p.add_run("(ПРОЕКТНО-ТЕХНОЛОГИЧЕСКУЮ) ПРАКТИКУ"), bold=True)

    details = [
        ("Тема", "Разработка информационно-поисковой системы для выбора теплообменных аппаратов"),
        ("Студент", "____________________________________________"),
        ("Группа", "____________"),
        ("Профильная организация", "СПбГТИ(ТУ), Санкт-Петербург"),
        ("Срок проведения", "06.07.2026–19.07.2026"),
        ("Срок сдачи отчёта", "19.07.2026"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Параметр"
    table.rows[0].cells[1].text = "Значение"
    for label, value in details:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    format_table(table, [["Параметр", "Значение"], *details])

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    set_run_font(p.add_run("Календарный план"), bold=True)
    plan = [
        ["Задача", "Срок"],
        ["Анализ литературных источников и интернет-ресурсов", "1–2 рабочий день"],
        ["Построение концептуальной и даталогической моделей БД", "3–4 рабочий день"],
        ["Построение UML-диаграмм USER и ADMIN", "5 рабочий день"],
        ["Разработка алгоритма фильтрации и ранжирования", "6–7 рабочий день"],
        ["Тестирование системы и оформление отчёта", "8–10 рабочий день"],
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for col, value in enumerate(plan[0]):
        table.rows[0].cells[col].text = value
    for row in plan[1:]:
        cells = table.add_row().cells
        for col, value in enumerate(row):
            cells[col].text = value
    format_table(table, plan)

    document.add_paragraph()
    for text in ("Задание принял к выполнению: ____________________", "Руководитель практики: ____________________"):
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        set_run_font(p.add_run(text), size=Pt(12))
    document.add_page_break()


def format_table(table, rows):
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, rows)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                shade_cell(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=Pt(11), bold=(row_index == 0))


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[name]
        style.font.name = BODY_FONT
        style.font.size = BODY_SIZE
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.page_break_before = False

    caption = document.styles["Caption"]
    caption.font.name = BODY_FONT
    caption.font.size = Pt(12)
    caption.font.bold = False
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_after = Pt(8)

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = BODY_FONT
        style.font.size = BODY_SIZE
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.paragraph_format.line_spacing = 1.5


def configure_sections(document):
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.header_distance = Cm(1)
        section.footer_distance = Cm(1)


def add_markdown_table(document, rows):
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for col, value in enumerate(rows[0]):
        table.rows[0].cells[col].text = value
    for values in rows[1:]:
        cells = table.add_row().cells
        for col, value in enumerate(values):
            cells[col].text = ""
            add_inline(cells[col].paragraphs[0], value, Pt(11))
    format_table(table, rows)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_table(lines, start):
    raw = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw.append(lines[index].strip())
        index += 1
    rows = []
    for line in raw:
        values = [value.strip() for value in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            continue
        rows.append(values)
    return rows, index


def add_source_content(document):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines = []
    first_h1 = True
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped == "```text" or stripped == "```powershell":
            in_code = True
            code_lines = []
            index += 1
            continue
        if stripped == "```" and in_code:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.left_indent = Cm(0.7)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(6)
            set_run_font(paragraph.add_run("\n".join(code_lines)), name=CODE_FONT, size=Pt(8))
            in_code = False
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped == "<!-- pagebreak -->":
            document.add_page_break()
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_markdown_table(document, rows)
            continue
        image_match = re.fullmatch(r"!\[(.+)]\((.+)\)", stripped)
        if image_match:
            caption, relative = image_match.groups()
            image_path = (SOURCE.parent / relative).resolve()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.keep_with_next = True
            from PIL import Image
            with Image.open(image_path) as image:
                aspect = image.width / image.height
            width_inches = min(6.35, 6.8 * aspect)
            picture = paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
            picture._inline.docPr.set("descr", caption)
            cap = document.add_paragraph(style="Caption")
            add_inline(cap, caption, Pt(12))
            index += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            paragraph = document.add_heading(level=level)
            if level == 1 and not first_h1:
                paragraph.paragraph_format.page_break_before = True
            add_inline(paragraph, text)
            if level == 1:
                first_h1 = False
                if text == "СОДЕРЖАНИЕ":
                    add_toc(document.add_paragraph())
            index += 1
            continue
        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
            index += 1
            continue
        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if number_match:
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, number_match.group(1))
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_inline(paragraph, stripped)
        index += 1


def build():
    document = Document()
    configure_styles(document)
    configure_sections(document)
    add_cover(document)
    add_assignment(document)
    add_source_content(document)
    configure_sections(document)

    section = document.sections[0]
    section.different_first_page_header_footer = True
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    add_page_field(footer_paragraph)
    section.first_page_footer.paragraphs[0].text = ""

    set_update_fields(document)
    document.core_properties.title = "Отчёт по учебной практике — ThermoSelect"
    document.core_properties.subject = "Информационно-поисковая система выбора теплообменных аппаратов"
    document.core_properties.author = "Студент СПбГТИ(ТУ)"
    document.core_properties.keywords = "ThermoSelect, теплообменники, учебная практика"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
